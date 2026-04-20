"""
Service: Resume Parser
Extracts text from PDF and DOCX files.
"""

import os
import fitz  # PyMuPDF
from docx import Document
from app.config import settings


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF"""
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        raise ValueError(f"Failed to parse PDF '{file_path}': {str(e)}")

    if not text.strip():
        raise ValueError(f"No text found in PDF '{file_path}'. It may be scanned/image-based.")

    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx"""
    text = ""
    try:
        doc = Document(file_path)

        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        # Extract from tables (resumes often have tables)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"

    except Exception as e:
        raise ValueError(f"Failed to parse DOCX '{file_path}': {str(e)}")

    if not text.strip():
        raise ValueError(f"No text found in DOCX '{file_path}'.")

    return text.strip()


def parse_resume(file_path: str) -> dict:
    """
    Parse a resume file and return extracted data.

    Args:
        file_path: Path to the resume file (PDF or DOCX)

    Returns:
        dict with keys: filename, file_path, raw_text, file_type, text_length
    """
    # Validate file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    # Get file extension
    filename = os.path.basename(file_path)
    ext = os.path.splitext(filename)[1].lower()

    # Validate extension
    if ext not in settings.ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file type '{ext}'. Allowed: {settings.ALLOWED_EXTENSIONS}")

    # Extract text based on file type
    if ext == ".pdf":
        raw_text = extract_text_from_pdf(file_path)
    elif ext == ".docx":
        raw_text = extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return {
        "filename": filename,
        "file_path": file_path,
        "raw_text": raw_text,
        "file_type": ext,
        "text_length": len(raw_text),
    }